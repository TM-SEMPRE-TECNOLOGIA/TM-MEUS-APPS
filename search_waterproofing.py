
import docx

doc_path = r'C:\Users\thiag\TM-MEUS-APPS\RELATÓRIO FOTOGRÁFICO - VILA TATETUBA - LEVANTAMENTO PREVENTIVO - ÁREA EXTERNA.docx'
doc = docx.Document(doc_path)

print("--- PESQUISA: IMPERMEABILIZAÇÃO ---")

# Busca em parágrafos
for i, p in enumerate(doc.paragraphs):
    if 'impermeabilização' in p.text.lower():
        print(f"P[{i}]: {p.text}")

# Busca em tabelas (onde costumam estar os itens do contrato)
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        row_text = " | ".join(c.text.strip().replace('\n', ' ') for c in row.cells)
        if 'impermeabilização' in row_text.lower():
             print(f"T[{t_idx}] R[{r_idx}]: {row_text}")
